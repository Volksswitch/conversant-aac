# -*- coding: utf-8 -*-
"""
Regenerates "AI-Driven AAC Product Overview.docx" — the plain-language,
non-technical overview of the project (the "From Communication to Conversation"
piece). Reconstructed June 2026 after the original generator was lost; styling
matches the original (Arial body, Word-default blue headings, Volksswitch footer
with page numbers). Run:  python generate-product-overview.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)     # title block
HEAD = RGBColor(0x2E, 0x74, 0xB5)     # Word-default heading blue
GREY = RGBColor(0x55, 0x55, 0x55)     # byline
FOOT = RGBColor(0x88, 0x88, 0x88)     # footer

doc = Document()

# --- base font: Arial throughout ---
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rpr.set(qn(a), "Arial")

# --- heading styles: Arial, default blue, original sizes, not bold ---
for sid, size in (("Heading 1", 16), ("Heading 2", 13)):
    st = doc.styles[sid]
    st.font.name = "Arial"
    st.font.size = Pt(size)
    st.font.bold = False
    st.font.color.rgb = HEAD


def title_line(text, size, color, bold=False, after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def body(text):
    return doc.add_paragraph(text)


# ============================ TITLE BLOCK ============================
title_line("AI-Driven AAC", 32, NAVY, bold=True, after=2)
title_line("From Communication to Conversation", 18, NAVY, after=8)
title_line("Kenneth R. Hackbarth | Volksswitch.org | June 2026", 11, GREY, after=14)

# ============================ CONTENT ============================
doc.add_heading("The Gap Nobody Talks About", level=1)
body("For decades, AAC devices have given non-speaking individuals a way to express needs, make choices, and share information. That is real progress. But there is a gap that rarely gets named directly:")
body("AAC devices support communication. Very few support conversation.")
body("These are not the same thing.")
body("Communication is the transfer of meaning from one person to another. A stop sign communicates. A fire alarm communicates. An AAC user selecting “I need help” communicates. The message goes out; the purpose is served.")
body("Conversation is something different. It is a joint activity — two or more people co-constructing meaning together, turn by turn, in real time. It is interactive, reciprocal, and time-sensitive. In conversation, each contribution responds to what came before and shapes what comes next.")
body("All conversation is communication. But not all communication is conversation.")
body("Most AAC systems are designed for communication. The result is that the people who use them are largely excluded from conversation — and from everything conversation makes possible.")

doc.add_heading("What AAC Users Actually Lose", level=1)
body("The loss of conversational access is often described as “slower speech.” That description misses the point entirely.")
body("When an AAC user cannot participate fluidly in conversation, they lose far more than speed. They lose access to the primary human mechanism for building identity, relationships, agency, humor, and belonging.")
body("Here is what the research shows non-conversational AAC users actually lose:")

doc.add_heading("Timing", level=2)
body("Conversation is deeply time-sensitive. Meaning is not carried only by words — it is carried by when something is said. A spoken person can say “Wait — that’s not what I meant” at the exact moment a misunderstanding begins. An AAC user may need many seconds to produce the same correction. By that time, the moment has passed.")
body("This affects jokes, corrections, emotional responses, quick questions, agreements, disagreements — any contribution that has to land at a specific moment to have its intended force. In conversation, late is often not merely late. Late can mean socially invisible.")

doc.add_heading("Turn-Taking Power", level=2)
body("Conversation depends on turn-taking — claiming the floor, holding it, yielding it, and redirecting it. AAC users are often forced into a passive structure: the partner asks, the AAC user answers. The partner guesses; the AAC user confirms or rejects.")
body("This turns conversation into interrogation. The AAC user loses the ability to initiate, to interrupt appropriately, to redirect a topic, to participate in fast back-and-forth exchange. Others control the structure of the interaction. The AAC user becomes a respondent rather than a co-participant.")

doc.add_heading("Spontaneity", level=2)
body("Much of conversation consists of small, unplanned, opportunistic moves: “That reminds me…” “Actually…” “Me too.” “That’s hilarious.” “I had the same problem.” These are not major messages. They are the small things that make conversation alive.")
body("Many AAC systems require deliberate message construction. That means the user may eventually be able to say something important — but cannot easily say the small things. They lose the ability to be casually present.")

doc.add_heading("Humor", level=2)
body("Humor depends on timing, surprise, shared background, and risk. An AAC user may be able to select “That’s funny” — but that is not the same as being able to make the joke. Humor is not ornamental. It is a major way people display intelligence, build intimacy, resist pity, and claim social equality. When a person cannot joke easily, others may treat them more solemnly, clinically, or childishly.")

doc.add_heading("Identity", level=2)
body("People are known through conversation. We become “the funny one,” “the thoughtful one,” “the skeptical one,” “the one who asks great questions.” When AAC only supports minimal functional communication, the person may instead be known mainly through needs: what they want, what they refuse, what hurts, what they can answer.")
body("That distorts identity. The person’s expressed self becomes smaller than the person’s actual self.")

doc.add_heading("The Summary", level=2)
body("When AAC supports communication but not conversation, users may lose: timing, agency, the ability to be known as complex and adult, the small exchanges that create friendship, the ability to explain and process emotions, the ability to intervene quickly in care decisions, access to private and confidential exchange, and the chance to develop thoughts with others rather than delivering finished messages.")
body("As one research synthesis puts it: the loss is not simply linguistic. It is dignity.")

doc.add_heading("The Root Cause — The Four-Second Problem", level=1)
body("Understanding why conversation has been out of reach for AAC users requires looking past the obvious answer.")
body("The obvious answer is speed. AAC users with significant motor limitations typically communicate at 10 to 20 words per minute. Comfortable conversation requires at least 80 words per minute. The gap is real.")
body("But speed alone is not the deepest barrier. The deepest barrier is the human aversion to awkward silence.")
body("Research shows that people become uncomfortable after as little as one to four seconds of silence in conversation. Repeated multi-second delays cause communication partners to interrupt, finish sentences for the AAC user, or disengage entirely.")
body("The result is that even a capable, articulate AAC user — one who has things to say and the device to say them — cannot hold a conversation partner’s attention long enough to be heard. They are confined to transactional exchanges: expressing needs and wants. The richer, interactional communication — the jokes, the debates, the small talk, the friendship-building — remains out of reach.")
body("This is not a matter of the AAC user’s capability. It is a technology gap.")

doc.add_heading("The Solution — AI-Generated Response Options", level=1)
body("Generative AI makes it possible to close that gap.")
body("Within seconds of a communication partner finishing a sentence, an AI system can generate several contextually appropriate, naturally worded response options for the AAC user to choose from. The user selects one with a single tap. The device speaks it immediately. The conversation continues at or near natural speed.")
body("The silence that used to strand the AAC user — and drive their partner away — is filled. Not with a generic placeholder, but with an actual response, ready for selection.")
body("This is the core idea. Everything else is refinement.")

doc.add_heading("Speaking As the User, Not For the User", level=2)
body("A critical distinction: the goal is not to speak for the AAC user. It is to speak as the AAC user.")
body("The AI does not generate generic responses that any person might give. It generates responses shaped by this specific person’s personality, opinions, relationships, and conversational goals. Over time, as the system learns who the user is, the responses sound more and more like them.")
body("The user retains final control over every word spoken. They select; the device speaks. Nothing is said without their choice.")

doc.add_heading("The User’s Communication Partner Needs Nothing", level=2)
body("The communication partner requires no device, no app, no training, and no special knowledge. From their perspective, they are simply having a conversation with someone who uses a device to communicate. The technology is entirely on the AAC user’s side.")

doc.add_heading("How It Works", level=1)
body("The core conversation loop has five steps:")
for step in [
    "The communication partner speaks. The system listens continuously.",
    "A text transcript of what the partner said appears on screen. The AAC user confirms the system heard correctly before proceeding. This step is non-negotiable — the user must be able to trust what the system heard.",
    "The AI generates a small set of response options. These appear on screen as large, tappable cards within seconds of the partner finishing their utterance. A brief, natural filler phrase (“Hmm, let me think…”) is spoken automatically to hold the conversational floor while options are being generated — so the silence never grows uncomfortable.",
    "The user selects a response with a single tap.",
    "The device speaks the selected response aloud. The conversation continues.",
]:
    doc.add_paragraph(step, style="List Number")
body("The entire loop — from partner finishing their utterance to the AAC user’s response being spoken — takes seconds. Not minutes.")

doc.add_heading("A Conversation Engine Built on How Talk Actually Works", level=2)
body("The system is grounded in Conversation Analysis, the empirical study of how real talk is structured. This means it understands more than turn-taking.")
body("It understands that conversations have structure — questions that expect answers, invitations that expect acceptance or decline, complaints that expect acknowledgment. It understands that a skilled conversationalist doesn’t only answer — they also initiate new directions, politely decline, request clarification, and gracefully close topics.")
body("The AI generates response options that represent these structurally distinct conversational moves — not three paraphrases of the same answer. The user can agree, decline with grace, push back, ask a follow-up, or signal that they didn’t understand — all from a single screen.")

doc.add_heading("What Makes This Different", level=1)
doc.add_heading("Designed for Conversation, Not Just Communication", level=2)
body("This system was built from the ground up around the goal of real-time conversational participation. The conversation loop, the response option design, the filler ladder, the timing — every decision was made in service of keeping a real conversation going.")
doc.add_heading("The System Learns Who You Are", level=2)
body("The system maintains a structured model of the user’s personality, interests, opinions, relationships, and communication goals. This is called the worldview model.")
body("Over time, as the user fills in more of their profile, responses sound more like them — their sense of humor, their way of expressing disagreement, their vocabulary, their characteristic concerns. The goal is a system that generates options so well-matched to this specific person that selecting one feels natural, not like settling.")
body("The worldview profile is built at the user’s own pace, a few questions at a time. No question is required. Even with a completely empty profile, the system is functional.")
doc.add_heading("Private by Design", level=2)
body("All user data — the worldview profile, conversation history, settings — stays on the user’s device. Nothing is sent to any external server except the AI API at the moment of generating responses. There is no cloud account holding the user’s personal information. The device is a Windows tablet; the data is a local folder.")
doc.add_heading("No Subscription. No Server Costs.", level=2)
body("Previous attempts to build AI-driven AAC systems have been shelved when funding ran out. A server-based architecture — where the project pays for computing on behalf of all users — is inherently fragile.")
body("This system avoids that failure mode entirely. It runs as a free web application with no backend server. Users create their own AI provider account and supply their own API key. The project incurs no ongoing costs. The user pays only for the conversations they actually have — typically a small fraction of a cent per exchange.")
body("This model scales to any number of users at near-zero cost to the project. It cannot be defunded.")
doc.add_heading("Free and Open Source", level=2)
body("The application is free to use. The source code is open source. No license, no subscription, no waiting list.")

doc.add_heading("Who It’s For Right Now", level=1)
body("The initial version is designed for:")
for item in [
    "Non-speaking individuals with cerebral palsy",
    "Literate (able to read response options on screen)",
    "Age 16 or older",
    "Using direct select (touch or pointer) as the primary access method",
    "Running a Windows tablet (Microsoft Surface is the recommended hardware)",
]:
    doc.add_paragraph(item, style="List Bullet")
body("The architecture is designed from the start to expand beyond this initial profile — to switch scanning, eye gaze, different literacy levels, and a broader age range. But the first version focuses here, where a working proof of concept delivers the most immediate value.")

doc.add_heading("What’s Coming", level=1)
doc.add_heading("Phase 1 — Core Conversation Loop (Available Now)", level=2)
body("The core loop described in this document — listen, transcript, generate options, select, speak — is built and working. This is the breakthrough: real-time conversational participation that was not possible with previous AAC systems.")
doc.add_heading("Phase 2 — Situational Awareness", level=2)
body("The system will add awareness of where the user is (via GPS), who they are speaking with (via face and voice recognition), and what’s on their calendar. Response options and conversation starters will become contextually appropriate to the specific setting and partner.")
doc.add_heading("Phase 3 — Full Worldview Shaping", level=2)
body("The complete worldview model — shaped by the user’s personality, values, relationships, and communication goals — will be integrated into every response generated. Phase 3 is where the system begins to speak truly as the user, not merely for them. Conversation history and a review loop allow the system to continuously improve its understanding of this specific person over time.")

doc.add_heading("Consider It for Your Clients and Family Members", level=1)
body("If you work with or care for a non-speaking individual who has the cognitive and literacy capacity to participate in conversation — and who is currently limited to transactional AAC use — this system is designed for them.")
body("It does not replace existing AAC vocabulary systems. It adds something those systems cannot provide: real-time conversational participation. The ability to respond, to initiate, to joke, to disagree, to be present in a conversation as themselves.")
body("For more information, contact Volksswitch.org.")

doc.add_heading("Glossary", level=1)
GLOSSARY = [
    ("AAC (Augmentative and Alternative Communication)", "A broad category of strategies, devices, and systems that supplement or replace natural speech for individuals who cannot speak or whose speech is difficult to understand. AAC ranges from low-tech picture boards and letter boards to high-tech speech-generating devices. This system is a high-tech AAC approach."),
    ("Adjacency Pair", "In Conversation Analysis: a two-part sequence in which one utterance (a question, greeting, invitation, or complaint) creates a structural expectation for a specific type of response. When someone asks a question, conversation norms make an answer expected — not silence. The conversation engine tracks open adjacency pairs to understand what kind of response is called for at any moment."),
    ("API Key", "A private credential that authorizes a software application to access an AI provider’s services over the internet. Users of this system create their own account with Anthropic (the AI provider) and supply their own API key. They are billed directly by Anthropic based only on the conversations they actually have — there is no subscription fee to the app itself."),
    ("Communication Partner", "The person speaking with the AAC user — the other participant in the conversation. This system places all technology on the AAC user’s side; the communication partner requires no device, no app, no training, and no special knowledge."),
    ("Conversation Analysis (CA)", "The academic discipline that studies how talk is structured in real social interaction — how people take turns, open and close topics, repair misunderstandings, and accomplish social actions through conversation. CA research (originating with Sacks, Schegloff, and Jefferson, 1974) provides the scientific foundation for this system’s conversation engine design."),
    ("Direct Select", "An access method in which the user directly touches or points to an item on a screen to select it, as opposed to indirect methods such as switch scanning (where a cursor moves through items automatically and the user activates a switch at the right moment) or eye-gaze systems (where gaze direction controls selection). The initial version of this system targets direct-select users."),
    ("Filler / Filler Phrase", "A brief spoken utterance (for example, “Hmm, let me think…” or “Just a second”) spoken automatically by the device to hold the conversational floor while AI response options are being generated. Fillers prevent the multi-second silence that causes communication partners to disengage. The system uses a tiered filler ladder: a short acknowledgment token within one second, followed by a longer projection phrase at two to three seconds if needed."),
    ("Generative AI", "A category of artificial intelligence that produces new content — text, audio, images, or other output — in response to a prompt. This system uses a generative AI language model to create contextually appropriate, naturally worded response options for the AAC user within seconds of the communication partner speaking."),
    ("LLM (Large Language Model)", "The type of AI model that generates response options. Large language models are trained on vast bodies of human text and can produce natural-sounding, contextually appropriate language in response to a prompt. Claude (developed by Anthropic) is the initial LLM used in this system, with support for additional providers planned."),
    ("Response Options", "The set of AI-generated conversational moves presented to the AAC user after the communication partner speaks. Each option represents a structurally distinct type of response — for example, agreement, graceful disagreement, a follow-up question, or a request for clarification. The user selects one with a single tap; the device speaks it aloud. Nothing is spoken without the user’s selection."),
    ("STT (Speech-to-Text)", "Automatic conversion of spoken audio into written text, also called transcription. The system uses STT to capture what the communication partner says and display it as a text transcript for the AAC user to confirm before response options are generated. Confirming the transcript is a non-negotiable step: the user must be able to verify the system heard correctly."),
    ("TTS (Text-to-Speech)", "Software that converts written text into spoken audio. The system uses TTS to speak the AAC user’s selected response aloud and to deliver filler phrases while options are being generated. The initial version uses the browser’s built-in TTS; voice banking (a personalized or cloned voice) is a planned future feature."),
    ("Turn-Taking", "The conversational mechanism by which participants alternate between speaking roles. Turn-taking involves claiming the floor, holding it, yielding it, and redirecting it — and doing so within narrow timing windows that most communication partners unconsciously enforce. AAC users are chronically disadvantaged by turn-taking timing requirements, which is the central problem this system is designed to solve."),
    ("Worldview Model", "A structured profile of the AAC user’s personality, values, interests, opinions, relationships, and communication goals. The worldview model is used to personalize AI-generated response options so they reflect who the user actually is — their sense of humor, their characteristic vocabulary, their way of expressing agreement or disagreement. The profile is built through a questionnaire completed entirely at the user’s own pace; no question is required, and the system is functional even with a completely empty profile."),
]
for term, definition in GLOSSARY:
    doc.add_heading(term, level=2)
    body(definition)

# ============================ FOOTER ============================
def add_field(run, instr):
    for t, attrs in [("begin", None), (None, instr), ("separate", None), (None, "1"), ("end", None)]:
        if t in ("begin", "separate", "end"):
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), t); run._r.append(fc)
        elif instr == attrs:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; run._r.append(it)
        else:
            tt = OxmlElement("w:t"); tt.text = attrs; run._r.append(tt)

sec = doc.sections[0]
footer = sec.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.text = ""
# top border on footer paragraph
pPr = fp._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top")
top.set(qn("w:val"), "single"); top.set(qn("w:color"), "CCCCCC")
top.set(qn("w:sz"), "4"); top.set(qn("w:space"), "4")
pBdr.append(top); pPr.append(pBdr)
# right tab stop at the right margin
tabs = OxmlElement("w:tabs")
tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "right")
right_pos = int((sec.page_width - sec.left_margin - sec.right_margin) / 635)  # EMU->twips
tab.set(qn("w:pos"), str(right_pos))
tabs.append(tab); pPr.append(tabs)


def foot_run(text=None, color=FOOT):
    r = fp.add_run(text or "")
    r.font.name = "Arial"; r.font.size = Pt(9)
    if color is not None:
        r.font.color.rgb = color
    return r

foot_run("Volksswitch.org | June 2026")
foot_run("\t", color=None)
add_field(foot_run(), "PAGE")

doc.save("AI-Driven AAC Product Overview.docx")
print("AI-Driven AAC Product Overview.docx generated.")
