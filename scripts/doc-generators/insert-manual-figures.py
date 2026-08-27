"""Put the two labeled figures into the Windows manual (Ken's comments 28 and 48).

  python insert-manual-figures.py

⚠ IDEMPOTENT BY CONSTRUCTION - it looks for its own captions and refuses if they are
already there. A generator that silently inserts a second copy every time it is run is
worse than one that fails, because nobody re-reads a 60-page manual to check.

The pictures go in as a run inside a new paragraph, which python-docx can only APPEND;
they are then moved into place with addnext on the XML. Insert positions are found by
matching the anchor's TEXT rather than by index, so an earlier edit cannot silently
shift a figure into the middle of another section.
"""
import os, sys, copy
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(HERE, '..', '..', 'Documents',
                   'Conversant AAC User Manual (Windows Chromebook Mac).docx')

FIGURES = [
    ('um-fig1-labeled.png',
     'Figure 1. The conversation screen, with the name of each region. These names are '
     'used throughout this manual.',
     'This section walks you through two short practice exchanges'),
    ('um-fig2-labeled.png',
     'Figure 2. The Composition Pane. It covers the Response Panel while it is open, and '
     'the on-screen keyboard takes the place of the Express Panel.',
     'Tap \u201cCancel\u201d at any time to close the Composition Pane without speaking.'),
]


def find(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    sys.exit('anchor not found: %s' % needle[:60])


def main():
    doc = docx.Document(DOC)
    if any('Figure 1.' in p.text for p in doc.paragraphs):
        sys.exit('REFUSING - the figures are already in this manual')

    for png, caption, anchor in FIGURES:
        target = find(doc, anchor)
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(os.path.join(HERE, png), width=Inches(6.4))
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(caption)
        run.italic = True
        # A figure and its caption must not be split across a page break, and the
        # caption must not be orphaned from what it describes.
        pic_p.paragraph_format.keep_with_next = True
        pic_p.paragraph_format.space_before = docx.shared.Pt(10)
        cap_p.paragraph_format.space_after = docx.shared.Pt(12)
        target._p.addnext(cap_p._p)
        target._p.addnext(pic_p._p)
        print('%s after "%s..."' % (png, anchor[:40]))

    # The pending one-word correction from the review: the module is called "How I
    # Sound" on screen. "Sounds Like Me" is the internal design-document name and a
    # user never sees it - exactly what the verify-against-source rule exists to catch.
    fixed = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if 'Sounds Like Me' in r.text:
                r.text = r.text.replace('Sounds Like Me', 'How I Sound')
                fixed += 1
    print('renamed Sounds Like Me -> How I Sound in %d run(s)' % fixed)
    doc.save(DOC)


main()
